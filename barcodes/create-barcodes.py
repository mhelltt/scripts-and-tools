import pandas as pd
import shutil
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from copy import deepcopy
from docx2pdf import convert
import os
import pathlib

df = pd.read_csv('skus-products.csv')

# slim_skus and classic_skus need to be defined within their functions, based on current product
# file permissions need to be copied and/or universal for the template

products = df['product'].unique()
slim_skus = df[df['sku'].str.contains('US')]['sku'].tolist()
classic_skus = df[~df['sku'].str.contains('US')]['sku'].tolist()
barcode_img_folder = os.path.abspath('barcode-img-files')+'/'
print(f'{len(products)} products in total')


def replaceText(table, search, sku):
    for row in table.rows:
        for cell in row.cells:
            if (cell.text == search):
                # remove placeholder search text
                cell.text = cell.text.replace(search, '')
                run = cell.paragraphs[0].add_run()
                img = f'{sku}.png'
                for root, dirs, files in os.walk(os.path.abspath('barcode-img-files')):
                    if img in files:
                        img_file = os.path.join(root, img)
                run.add_picture(
                    img_file, width=Inches(3.5))
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_slims(product):
    slims_temp = 'barcode_temp.docx'
    slims_temp_doc = f'{pathlib.Path(__file__).parent.resolve()}/barcode-pdfs/{slims_temp}'
    shutil.copyfile(
        'template.docx', slims_temp_doc)
    document = Document(slims_temp_doc)
    product_slims = df[(df['sku'].str.contains('US')) & (
        df['product'] == product)]['sku'].tolist()
    print(f'Writing {product} slims...')
    for sku in product_slims:
        temp_tables = document.tables
        tbl = temp_tables[len(temp_tables)-1]._tbl
        tbl_obj = temp_tables[len(temp_tables)-1]
        new_tbl = deepcopy(tbl)
        replaceText(tbl_obj, '%BARCODE%', sku)
        if (len(temp_tables) < len(product_slims)+1):
            paragraph = document.add_paragraph()
            paragraph._p.addnext(new_tbl)
            paragraph.add_run().add_break(WD_BREAK.PAGE)
    document.save(slims_temp_doc)
    slims_pdf = f'{pathlib.Path(__file__).parent.resolve()}/barcode-pdfs/{product}_slims.pdf'
    print(f'Converting {product} slims doc to pdf')
    convert(slims_temp_doc, slims_pdf)
    os.remove(slims_temp_doc)
    print(f'{product} slims pdf ready')


def create_classics(product):
    classics_temp = 'barcode_temp.docx'
    classics_temp_doc = f'{pathlib.Path(__file__).parent.resolve()}/barcode-pdfs/{classics_temp}'
    shutil.copyfile(
        'template.docx', classics_temp_doc)
    document = Document(classics_temp_doc)
    product_classics = df[(~df['sku'].str.contains('US')) & (
        df['product'] == product)]['sku'].tolist()
    print(f'Writing {product} classics...')
    for sku in product_classics:
        temp_tables = document.tables
        tbl = temp_tables[len(temp_tables)-1]._tbl
        tbl_obj = temp_tables[len(temp_tables)-1]
        new_tbl = deepcopy(tbl)
        replaceText(tbl_obj, '%BARCODE%', sku)
        if (len(temp_tables) < len(product_classics)+1):
            paragraph = document.add_paragraph()
            paragraph._p.addnext(new_tbl)
            paragraph.add_run().add_break(WD_BREAK.PAGE)
    document.save(classics_temp_doc)
    classics_pdf = f'{pathlib.Path(__file__).parent.resolve()}/barcode-pdfs/{product}_classics.pdf'
    print(f'Converting {product} classics doc to pdf')
    convert(classics_temp_doc, classics_pdf)
    os.remove(classics_temp_doc)
    print(f'{product} classics pdf ready')


if __name__ == "__main__":
    for product in products:
        create_slims(product)
        create_classics(product)
