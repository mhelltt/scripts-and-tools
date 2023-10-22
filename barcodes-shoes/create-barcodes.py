import pandas as pd
import shutil
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from copy import deepcopy
from docx2pdf import convert
import os
import pathlib

df = pd.read_csv('skus-products.csv')

# file permissions need to be copied and/or universal for the template

products = df['product'].unique()
skus = df['sku'].tolist()
barcode_img_folder = os.path.abspath('barcode-img-files')+'/'
print(f'{len(products)} products in total')
                
def addInfo(table, style, size, color):
    for row in table.rows:
        for cell in row.cells:
            p = cell.add_paragraph()
            r = p.add_run()
            r.add_break()
            r.add_text(style)
            r.add_break()
            r.add_text(size)
            r.add_break()
            r.add_text(color)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.left_indent = Inches(0.1)
            f = r.font
            f.size = Pt(8)


def addBarcode(table, sku):
    for row in table.rows:
        for cell in row.cells:
            run = cell.add_paragraph().add_run()
            img = f'{sku}.png'
            for root, dirs, files in os.walk(os.path.abspath('barcode-img-files')):
                if img in files:
                    img_file = os.path.join(root, img)
            run.add_picture(
                    img_file, width=Inches(2.5))

def create_stickers(product):
    stickers_temp = 'barcode_temp.docx'
    stickers_temp_doc = f'{pathlib.Path(__file__).parent.resolve()}/barcode-pdfs/{stickers_temp}'
    shutil.copyfile(
        'template.docx', stickers_temp_doc)
    document = Document(stickers_temp_doc)
    product_skus = df[(df['product'] == product)]['sku'].tolist()
    print(f'Writing {product} ...')
    for sku in product_skus:
        sku_style = df.loc[(df['sku'] == sku)]['style'].values[0]
        sku_color = df.loc[(df['sku'] == sku)]['color'].values[0]
        sku_size = df.loc[(df['sku'] == sku)]['size'].values[0]
        
        temp_tables = document.tables
        tbl = temp_tables[len(temp_tables)-1]._tbl
        tbl_obj = temp_tables[len(temp_tables)-1]
        new_tbl = deepcopy(tbl)
        
        style_info = f'Style: {str(sku_style)}'
        color_info = f'Color: {str(sku_color)}'
        size_info = f'Size: {str(sku_size)}'

        addInfo(tbl_obj, style_info, color_info, size_info)
        addBarcode(tbl_obj, sku)
        
        if (len(temp_tables) < len(product_skus)+1):
            paragraph = document.add_paragraph()
            paragraph._p.addnext(new_tbl)
            paragraph.add_run().add_break(WD_BREAK.PAGE)
    document.save(stickers_temp_doc)
    stickers_pdf = f'{pathlib.Path(__file__).parent.resolve()}/barcode-pdfs/{product}.pdf'
    print(f'Converting {product} doc to pdf')
    convert(stickers_temp_doc, stickers_pdf)
    os.remove(stickers_temp_doc)
    print(f'{product} pdf ready')

if __name__ == "__main__":
    for product in products:
        create_stickers(product)
